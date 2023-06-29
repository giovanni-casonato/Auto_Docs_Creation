import os
from docx import Document
import pandas as pd

# Structure of the Document
file_structure = """{Name} was approved by the Loan Committee on {Status Date} at 2:30 PM."""

# Find and read Excel file with the Data we need
excel_file_path = "C:/Users/giovanni.casonato/Desktop/Pycharm/AutoDoc Creation/Brokers_Names.xlsx
df = pd.read_excel(excel_file_path)

# Convert the Excel data into a dictionary
brokers_data = df.to_dict(orient="records")
file_data = []
Names_and_Dates = ["Name", "Status Date"]

# Select only the values that have Relationship Status as Active
# Save the variables for the keys 'Name' and 'Relationship Status'
for values in brokers_data:
    if values['Relationship Status'] == 'Active':
        data = {key: values[key] for key in Names_and_Dates}

        # Remove the time 00:00:00 portion from the "Status Date" key and add Keys and Values to file_data
        data['Status Date'] = values['Status Date'].date().strftime('%Y-%m-%d')
        file_data.append(data)

# Select a folder to save the docs
output_folder = "C:/Users/giovanni.casonato/Desktop/Approval Docs/"

# The loop returns an iterable that provides tuples with the index and the element
# The values in the 'Name' key is used in the heading of the file
# The Data from the file_data is saved in the structure variable
for i, data in enumerate(file_data):
    doc = Document()
    doc.add_heading(f"Approval Form for {data['Name']}", level=2)
    doc.add_paragraph(file_structure.format(**data))
# The file name is cleaned from spaces and replace to _
    name = str(data['Name'])
    filename = os.path.join(output_folder, f"Approval_{name.replace(' ', '_')}.docx")
    doc.save(filename)
    print(f"Created file: {filename}")


